import re
from pathlib import Path
from typing import Any

from pydantic import field_validator

from ..models import StrictModel
from ..source_types import SourceCommand, SourceContext


class DocxExtractTextInput(StrictModel):
    """Input for docx.extract_text.

    Attributes:
        path: Workspace-relative DOCX path.
    """

    path: str

    @field_validator("path")
    @classmethod
    def require_docx_extension(cls, value: str) -> str:
        """Requires the path to target a DOCX file.

        Args:
            value: Workspace-relative path from verifier.json.

        Returns:
            The unchanged path when it has a DOCX extension.

        Raises:
            ValueError: If the path does not end in ``.docx``.
        """
        if Path(value).suffix.lower() != ".docx":
            raise ValueError("docx.extract_text requires a .docx file")
        return value


class DocxExtractTextOutput(StrictModel):
    """Output from docx.extract_text.

    Attributes:
        text: Extracted body paragraph and table text.
    """

    text: str


class ExtractText(SourceCommand[DocxExtractTextInput, DocxExtractTextOutput]):
    """Extracts text from DOCX body paragraphs and tables."""

    name = "extract_text"
    input_model = DocxExtractTextInput
    output_model = DocxExtractTextOutput

    def run(self, source_input: DocxExtractTextInput, context: SourceContext) -> DocxExtractTextOutput:
        """Runs DOCX text extraction.

        Args:
            source_input: Validated DOCX extraction input.
            context: Source runtime context.

        Returns:
            Extracted DOCX text.

        """
        resolved = context.resolve_path(source_input.path)
        return DocxExtractTextOutput(text=extract_docx(resolved, context.max_content_chars))


class DocxInspectDocumentInput(StrictModel):
    """Input for ``docx.inspect_document``.

    Attributes:
        path: Workspace-relative DOCX path.
    """

    path: str

    @field_validator("path")
    @classmethod
    def require_docx_extension(cls, value: str) -> str:
        """Requires the path to target a DOCX file.

        Args:
            value: Workspace-relative path from verifier.json.

        Returns:
            The unchanged path when it has a DOCX extension.

        Raises:
            ValueError: If the path does not end in ``.docx``.
        """
        if Path(value).suffix.lower() != ".docx":
            raise ValueError("docx.inspect_document requires a .docx file")
        return value


class DocumentHeading(StrictModel):
    """One heading paragraph with its outline level.

    Attributes:
        text: The heading's visible text.
        level: The trailing number of its ``Heading N`` style, or ``0`` for a
            heading style carrying no level.
    """

    text: str
    level: int


class DocxInspectDocumentOutput(StrictModel):
    """Document structure expressed as counts and heading text.

    Attributes:
        paragraph_count: Body paragraphs holding visible text. Empty spacing
            paragraphs and the paragraph an inline image sits in are excluded,
            so "at least five paragraphs" grades what a reader would count.
        word_count: Whitespace-separated words across body paragraphs and table
            cells, so a document whose content lives in a table is not reported
            as wordless.
        table_count: Body tables.
        section_count: Word sections, which carry page setup and headers.
        heading_count: Paragraphs styled ``Heading *``. Exact even when
            ``heading_texts`` is truncated.
        heading_texts: Heading text as a flat list, capped to the configured
            source content limit. Flat on purpose: ``$.heading_texts contains
            "Recommendation"`` is a single JSONPath match holding a list, while
            a multi-match path over ``headings[*].text`` changes the comparison
            semantics ``compare.actual_from_matches`` warns about.
        headings: The same headings with their outline levels, in document
            order, truncated alongside ``heading_texts``.
        image_count: Inline images.
        has_header: Whether any section header carries visible text.
        has_footer: Whether any section footer carries visible text.
    """

    paragraph_count: int
    word_count: int
    table_count: int
    section_count: int
    heading_count: int
    heading_texts: list[str]
    headings: list[DocumentHeading]
    image_count: int
    has_header: bool
    has_footer: bool


class InspectDocument(
    SourceCommand[DocxInspectDocumentInput, DocxInspectDocumentOutput]
):
    """Report headings, counts, images, and header/footer presence."""

    name = "inspect_document"
    input_model = DocxInspectDocumentInput
    output_model = DocxInspectDocumentOutput

    def run(
        self, source_input: DocxInspectDocumentInput, context: SourceContext
    ) -> DocxInspectDocumentOutput:
        """Runs DOCX structure inspection.

        Args:
            source_input: Validated document inspection input.
            context: Source runtime context.

        Returns:
            Document structure as counts and bounded heading text.
        """
        resolved = context.resolve_path(source_input.path)
        return inspect_docx_document(resolved, context.max_content_chars)


def inspect_docx_document(path: Path, limit: int) -> DocxInspectDocumentOutput:
    """Reads document structure and heading text.

    Args:
        path: Resolved DOCX file path.
        limit: Maximum number of heading characters to return.

    Returns:
        Document structure, with heading evidence capped to ``limit``.
    """
    from docx import Document

    doc = Document(str(path))
    paragraph_count = 0
    word_count = 0
    heading_count = 0
    heading_texts: list[str] = []
    headings: list[DocumentHeading] = []
    heading_total = 0

    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text:
            paragraph_count += 1
            word_count += len(text.split())
        style_name = str(getattr(paragraph.style, "name", "") or "")
        if not style_name.startswith("Heading"):
            continue
        heading_count += 1
        if heading_total >= limit:
            continue
        heading_total = append_limited(heading_texts, text, heading_total, limit)
        headings.append(DocumentHeading(text=text, level=heading_level(style_name)))

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                word_count += len(cell.text.split())

    return DocxInspectDocumentOutput(
        paragraph_count=paragraph_count,
        word_count=word_count,
        table_count=len(doc.tables),
        section_count=len(doc.sections),
        heading_count=heading_count,
        heading_texts=heading_texts,
        headings=headings,
        image_count=len(doc.inline_shapes),
        has_header=any(has_visible_text(section.header) for section in doc.sections),
        has_footer=any(has_visible_text(section.footer) for section in doc.sections),
    )


def heading_level(style_name: str) -> int:
    """Reads the outline level from a ``Heading N`` style name.

    Args:
        style_name: Paragraph style name.

    Returns:
        The trailing number, or ``0`` for a heading style carrying none.
    """
    match = re.search(r"(\d+)\s*$", style_name)
    return int(match.group(1)) if match else 0


def has_visible_text(header_footer: Any) -> bool:
    """Whether a section header or footer carries visible text.

    A Word section always has a header and a footer object, so presence proves
    nothing; a letterhead or a page number is what the prompt actually means.

    Args:
        header_footer: python-docx header or footer object.

    Returns:
        True when any of its paragraphs holds non-whitespace text.
    """
    return any(paragraph.text.strip() for paragraph in header_footer.paragraphs)


def extract_docx(path: Path, limit: int) -> str:
    """Extracts body paragraph and table text from a DOCX file.

    Args:
        path: Resolved DOCX file path.
        limit: Maximum number of characters to return.

    Returns:
        Extracted text capped to ``limit`` characters.
    """
    from docx import Document

    doc = Document(str(path))
    parts: list[str] = []
    total = 0
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            total = append_limited(parts, paragraph.text, total, limit)
            if total >= limit:
                return "\n".join(parts)[:limit]
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text and cell.text.strip()]
            if cells:
                total = append_limited(parts, " | ".join(cells), total, limit)
                if total >= limit:
                    return "\n".join(parts)[:limit]
    # Headers/footers (letterheads, page numbers) live outside doc.paragraphs.
    for section in doc.sections:
        for header_footer in (section.header, section.footer):
            for paragraph in header_footer.paragraphs:
                if paragraph.text.strip():
                    total = append_limited(parts, paragraph.text, total, limit)
                    if total >= limit:
                        return "\n".join(parts)[:limit]
    return "\n".join(parts)[:limit]


def append_limited(parts: list[str], text: str, total: int, limit: int) -> int:
    """Appends text and returns the updated character count.

    Args:
        parts: Text fragments collected so far.
        text: Text fragment to append.
        total: Current approximate character count.
        limit: Maximum desired character count.

    Returns:
        Updated approximate character count.
    """
    parts.append(text)
    return total + len(text) + 1


COMMANDS = (ExtractText(), InspectDocument())
